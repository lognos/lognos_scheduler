"""Service for processing email attachments (PDF, Word, Excel)."""

import logfire
from typing import Literal
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
import os
import io
import pandas as pd
import docx
from openpyxl import load_workbook

from backend.config.settings import settings
from backend.models.attachment import (
    AttachmentSummary,
    AttachmentMetadata,
    EmailAttachmentInfo,
)

class ExtractionResult(BaseModel):
    """Structured output from Gemini extraction."""
    summary: str = Field(description="Concise summary of the document content")
    extracted_text: str = Field(description="Key text extracted from the document")
    page_count: int | None = Field(None, description="Estimated page count if applicable")

class AttachmentProcessorService:
    """
    Processes email attachments and extracts text content using Gemini.
    """
    
    MAX_TEXT_LENGTH = 10000  # chars
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_TOTAL_SIZE = 25 * 1024 * 1024  # 25MB per email
    
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
    }
    
    def __init__(self):
        """Initialize Gemini client singleton."""
        api_key = settings.google_api_key or settings.gemini_api_key
        if not api_key:
            logfire.warn("Google API key not found, attachment processing will fail")
            self.client = None
        else:
            self.client = genai.Client(api_key=api_key)

    @logfire.instrument("AttachmentProcessor.process_attachment")
    async def process_attachment(
        self,
        attachment: EmailAttachmentInfo,
        context: str | None = None,
    ) -> AttachmentSummary:
        """
        Process a single attachment using Gemini Native Document Understanding.
        
        Args:
            attachment: Attachment info with content bytes
            
        Returns:
            AttachmentSummary with extracted text and metadata
        """
        if not attachment.content_bytes:
            return self._create_error_summary(attachment, "No content bytes provided")
            
        if attachment.size_bytes > self.MAX_FILE_SIZE:
             return self._create_error_summary(attachment, f"File size exceeds limit ({self.MAX_FILE_SIZE/1024/1024}MB)")

        if attachment.content_type not in self.SUPPORTED_TYPES:
             return self._create_error_summary(
                 attachment, 
                 f"Unsupported file type: {attachment.content_type}. (Supported: {', '.join(self.SUPPORTED_TYPES.keys())})",
                 status="partial"
             )

        if not self.client:
             return self._create_error_summary(attachment, "Gemini client not initialized")

        # Generate content with structured output
        if attachment.content_type == "application/pdf":
            return await self._process_pdf_native(attachment, context)
        elif attachment.content_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
            return await self._process_excel_local(attachment, context)
        elif attachment.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._process_word_local(attachment, context)
        
        # Fallback (should be caught by supported types check, but for safety)
        return self._create_error_summary(attachment, "Unsupported type for processing logic")

    async def _process_pdf_native(
        self,
        attachment: EmailAttachmentInfo,
        context: str | None = None,
    ) -> AttachmentSummary:
        try:
             # Construct system prompt
            system_instruction = (
                "Analyze this document. Provide a concise summary of key topics, risks, and actions. "
                "Also extract the most important text content, up to roughly 2000 words."
            )
            if context:
                system_instruction += f" Keep this specific goal in mind: {context}"

            response = await self.client.aio.models.generate_content(
                model=settings.doc_processor_model_name,
                contents=[
                    types.Part.from_bytes(
                        data=attachment.content_bytes,
                        mime_type=attachment.content_type
                    ),
                    system_instruction
                ],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=ExtractionResult,
                )
            )
            
            # Parse structured output
            if not response.parsed:
                 return self._create_error_summary(attachment, "Failed to parse Gemini response")
                 
            result: ExtractionResult = response.parsed
            
            # Truncate text if needed
            final_text = result.extracted_text[:self.MAX_TEXT_LENGTH]
            if len(result.extracted_text) > self.MAX_TEXT_LENGTH:
                final_text += "...[truncated]"

            logfire.info(
                "Attachment processed successfully",
                filename=attachment.filename,
                text_length=len(final_text)
            )
            
            return AttachmentSummary(
                metadata=AttachmentMetadata(
                    filename=attachment.filename,
                    content_type=attachment.content_type,
                    size_bytes=attachment.size_bytes,
                    page_count=result.page_count
                ),
                extracted_text=final_text,
                summary=result.summary,
                processing_status="success",
            )
            
        except Exception as e:
            logfire.error(f"Attachment processing failed: {e}", filename=attachment.filename)
            return self._create_error_summary(attachment, str(e), status="failed")

    async def _process_excel_local(
        self,
        attachment: EmailAttachmentInfo,
        context: str | None = None,
    ) -> AttachmentSummary:
        try:
            # 1. Filter Visible Sheets using openpyxl
            visible_sheets = []
            with io.BytesIO(attachment.content_bytes) as bio:
                wb = load_workbook(bio, read_only=True, data_only=True)
                for sheet in wb.worksheets:
                    if sheet.sheet_state == 'visible':
                        visible_sheets.append(sheet.title)
                wb.close()
            
            if not visible_sheets:
                return self._create_error_summary(attachment, "No visible sheets found", status="partial")

            # 2. Parse Data (Visible Sheets Only)
            with io.BytesIO(attachment.content_bytes) as bio:
                # Only read specific sheets to save memory
                dfs = pd.read_excel(
                    bio, 
                    sheet_name=visible_sheets, 
                    engine='openpyxl', 
                    nrows=100,  # Cap rows early for efficiency
                    header=None  # Treat first row as data, not header
                )

            # 3. Build LLM Context
            context_parts = []
            context_parts.append(f"Document: {attachment.filename}")
            if context:
                context_parts.append(f"User Analysis Instruction: {context}")
            
            context_parts.append(f"Contains {len(visible_sheets)} visible sheets: {', '.join(visible_sheets)}")

            # Construct Markdown tables
            current_char_count = 0
            MAX_CONTEXT_CHARS = 20000 
            
            for name, df in dfs.items():
                if current_char_count >= MAX_CONTEXT_CHARS:
                    context_parts.append(f"\n[Sheet '{name}' omitted due to size limits]")
                    continue
                    
                # Clean generic unnamed columns if needed
                df = df.dropna(how='all', axis=0).dropna(how='all', axis=1)
                
                rows, cols = df.shape
                display_df = df.head(50) # Show top 50 rows
                markdown_table = display_df.to_markdown(index=False)
                
                sheet_section = f"\n## Sheet: {name} (Snippet)\n{markdown_table}\n"
                context_parts.append(sheet_section)
                current_char_count += len(sheet_section)

            full_context_text = "\n".join(context_parts)

            # 4. Standardized Gemini Prompt
            system_instruction = (
                "Analyze this Excel data. "
                "Ignore hidden sheets or unavailable data. "
                "CRITICAL: Do not invent or hallucinate data. If the file appears empty or has no relevant data, "
                "explicitly state that in the summary."
            )
            if context:
                system_instruction += f" Focus specifically on this goal: {context}"
            else:
                system_instruction += " Summarize key data points, dates, and status columns."

            # 5. Call Gemini
            response = await self.client.aio.models.generate_content(
                model=settings.doc_processor_model_name,
                contents=[full_context_text, system_instruction],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=ExtractionResult,
                    temperature=0.0,  # Strict determinism
                )
            )
            
            if not response.parsed:
                 return self._create_error_summary(attachment, "Failed to parse Gemini response")
                 
            result = response.parsed
            
            # Truncate text if needed
            final_text = result.extracted_text[:self.MAX_TEXT_LENGTH]
            if len(result.extracted_text) > self.MAX_TEXT_LENGTH:
                final_text += "...[truncated]"

            logfire.info(
                "Excel Attachment processed successfully",
                filename=attachment.filename,
                text_length=len(final_text)
            )
            
            return AttachmentSummary(
                metadata=AttachmentMetadata(
                    filename=attachment.filename,
                    content_type=attachment.content_type,
                    size_bytes=attachment.size_bytes,
                    sheet_count=len(visible_sheets),
                    sheet_names=visible_sheets
                ),
                extracted_text=final_text,
                summary=result.summary,
                processing_status="success",
            )
            
        except Exception as e:
            lcf.error(f"Excel processing failed: {e}", filename=attachment.filename)
            return self._create_error_summary(attachment, str(e), status="failed")

    async def _process_word_local(
        self,
        attachment: EmailAttachmentInfo,
        context: str | None = None,
    ) -> AttachmentSummary:
        try:
            with io.BytesIO(attachment.content_bytes) as bio:
                doc = docx.Document(bio)
                
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Simple table text extraction
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        full_text.append(" | ".join(row_data))

            joined_text = "\n\n".join(full_text)
            paragraph_count = len(full_text)

            # Check hallucination risk (Empty file)
            if not joined_text.strip():
                 return self._create_error_summary(attachment, "Document appears empty of text", status="partial")

            # Truncate if needed
            final_text = joined_text
            if len(joined_text) > self.MAX_TEXT_LENGTH:
                 final_text = joined_text[:self.MAX_TEXT_LENGTH] + "...[truncated]"

            # Build Gemini Context
            system_instruction = (
                "Analyze this Word document content. "
                "CRITICAL: Do not invent or hallucinate data. If the content is missing or insufficient, explicitly state that."
            )
            if context:
                system_instruction += f" Focus specifically on this goal: {context}"

            response = await self.client.aio.models.generate_content(
                model=settings.doc_processor_model_name,
                contents=[final_text, system_instruction],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=ExtractionResult,
                    temperature=0.0,
                )
            )

            if not response.parsed:
                 return self._create_error_summary(attachment, "Failed to parse Gemini response")
                 
            result = response.parsed

            logfire.info(
                "Word Attachment processed successfully",
                filename=attachment.filename,
                text_length=len(final_text)
            )

            return AttachmentSummary(
                metadata=AttachmentMetadata(
                    filename=attachment.filename,
                    content_type=attachment.content_type,
                    size_bytes=attachment.size_bytes,
                    paragraph_count=paragraph_count,
                ),
                extracted_text=final_text[:self.MAX_TEXT_LENGTH], # Ensure stored text is also truncated
                summary=result.summary,
                processing_status="success",
            )

        except Exception as e:
            logfire.error(f"Word processing failed: {e}", filename=attachment.filename)
            return self._create_error_summary(attachment, f"Word parse error: {str(e)}", status="failed")

    def _create_error_summary(
        self, 
        attachment: EmailAttachmentInfo, 
        error: str, 
        status: Literal["failed", "partial"] = "failed"
    ) -> AttachmentSummary:
        """Helper to create error summaries."""
        return AttachmentSummary(
            metadata=AttachmentMetadata(
                filename=attachment.filename,
                content_type=attachment.content_type,
                size_bytes=attachment.size_bytes,
            ),
            extracted_text="",
            summary=f"Processing failed: {error}",
            processing_status=status,
            error_message=error,
        )

    async def process_all_attachments(
        self,
        attachments: list[EmailAttachmentInfo],
        context: str | None = None,
    ) -> list[AttachmentSummary]:
        """Process multiple attachments with total size check."""
        summaries = []
        total_size = 0
        
        for attachment in attachments:
            # Check total size limit
            if total_size + attachment.size_bytes > self.MAX_TOTAL_SIZE:
                logfire.warn(
                    "Total attachment size limit exceeded, skipping remaining",
                    current_total=total_size,
                    limit=self.MAX_TOTAL_SIZE
                )
                summaries.append(self._create_error_summary(
                    attachment, 
                    "Total email attachment size limit exceeded", 
                    status="partial"
                ))
                continue

            summary = await self.process_attachment(attachment, context)
            summaries.append(summary)
            if summary.processing_status == "success":
                total_size += attachment.size_bytes
                
        return summaries
